# report_generator_docx.py

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
# from docx.enum.section import WD_ORIENT # Не используется
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
# Убедитесь, что все необходимые модели импортированы
from .models import db, Publication, PublicationAuthor, PublicationTypeDisplayName, User, PublicationType
from sqlalchemy.orm import joinedload, selectinload
from datetime import datetime, timedelta
from typing import Optional, List
import logging

logger = logging.getLogger(__name__)

# --- Функция set_cell_border (без изменений) ---
def set_cell_border(cell, **kwargs):
    """
    Helper function to set borders on a table cell.
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000", "space": "0"})
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None: tcBorders = OxmlElement("w:tcBorders"); tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "start", "end", "insideH", "insideV"):
        if edge in kwargs:
            tag = f"w:{edge}"; border_element = tcBorders.find(qn(tag))
            if border_element is None: border_element = OxmlElement(tag); tcBorders.append(border_element)
            for k, v in kwargs[edge].items(): border_element.set(qn(f"w:{k}"), str(v))
# --- ---

# --- Вспомогательная функция для заголовка секции (без изменений) ---
def _add_section_header(table, text, num_columns):
    """ Вспомогательная функция для добавления строки-заголовка секции ('а)', 'б)', 'в)') """
    row = table.add_row()
    merged_cell = row.cells[0].merge(row.cells[num_columns - 1])
    p = merged_cell.paragraphs[0]
    run = p.add_run(text); run.font.name = 'Times New Roman'; run.font.size = Pt(11); run.font.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # Границы можно не ставить, если стиль таблицы их задает
    # border_args = {"sz": 4, "val": "single", "color": "000000"}
    # set_cell_border(merged_cell, top=border_args, bottom=border_args, start=border_args, end=border_args)

# --- ОСНОВНАЯ ФУНКЦИЯ С ИЗМЕНЕНИЯМИ v5 ---
def generate_docx_report(user_id: int, start_date: Optional[datetime] = None, end_date: Optional[datetime] = None) -> bytes:
    """
    Генерирует DOCX отчет (формат ПМИ) по публикациям пользователя v5.
    - Строгая группировка: б) только ПЭВМ, в) только Уч.пособие, а) остальные.
    - Отображает типы с большой буквы, ПЭВМ - капсом.
    - Корректно формирует выходные данные для каждой секции.
    - В блоке 'а' всегда добавляются Место/Издательство, если есть.
    - Увеличена ширина столбца "Форма работы".
    - Форматирует объем без лишних нулей.
    """
    logger.info(f"Генерация DOCX отчета (ПМИ v5) для user_id={user_id}, start={start_date}, end={end_date}")

    # 1. Загрузка данных пользователя и публикаций
    user = User.query.get(user_id)
    if not user:
        logger.error(f"Пользователь с ID {user_id} не найден."); document_empty = Document(); document_empty.add_paragraph(f"Ошибка: Пользователь с ID {user_id} не найден.")
        buffer = BytesIO(); document_empty.save(buffer); buffer.seek(0); return buffer.getvalue()

    query = Publication.query.filter(Publication.user_id == user_id, Publication.status == 'published')\
        .options(selectinload(Publication.authors), joinedload(Publication.type), joinedload(Publication.display_name))
    if start_date: query = query.filter(Publication.published_at >= start_date)
    if end_date: query = query.filter(Publication.published_at < (end_date + timedelta(days=1)))
    publications = query.order_by(Publication.year.desc(), Publication.title).all()
    logger.info(f"Найдено публикаций до группировки: {len(publications)}")

    # 2. Группировка публикаций (СТРОГАЯ)
    group_a, group_b, group_c = [], [], []
    pvm_type_name = "ppc" # Внутреннее имя типа для ПЭВМ
    uchebnoe_posobie_display_name = "Учебное пособие" # Отображаемое имя для группы 'в'
    monograph_display_name = "Монография" # Используется для определения формата страниц

    for pub in publications:
        pub_type_name = pub.type.name if pub.type else None
        pub_display_name = pub.display_name.display_name if pub.display_name else None

        if pub_type_name == pvm_type_name:
            group_b.append(pub) # Группа б) - Строго ПЭВМ
        elif pub_display_name == uchebnoe_posobie_display_name:
            group_c.append(pub) # Группа в) - Строго Учебное пособие
        else:
            group_a.append(pub) # Все остальное в группу а)

    logger.info(f"Строгая группировка: А={len(group_a)}, Б={len(group_b)}, В={len(group_c)}")

    # 3. Создание документа и настройка
    document = Document(); section = document.sections[0]
    section.page_height = Cm(29.7); section.page_width = Cm(21.0); section.left_margin = Cm(2.0); section.right_margin = Cm(1.0); section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)
    style = document.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(11); style.paragraph_format.space_after = Pt(0); style.paragraph_format.line_spacing = 1.0

    # 4. Заголовок документа
    p1 = document.add_paragraph(); p1.add_run('Список').bold = True; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = document.add_paragraph(); user_position = getattr(user, 'position', "(Должность)"); department_name = getattr(user, 'department', "кафедры ПМИ")
    p2.add_run(f"научных и учебно-методических работ {user_position} {department_name}"); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3 = document.add_paragraph(); p3.add_run(user.full_name).bold = True; p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if start_date or end_date:
        date_str = "за период"
        if start_date: date_str += f" с {start_date.strftime('%d.%m.%Y')}";
        if end_date: date_str += f" по {end_date.strftime('%d.%m.%Y')}"
        p_date = document.add_paragraph(); p_date.add_run(date_str); p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    document.add_paragraph() # Пустая строка

    # 5. Основная таблица
    if not publications: document.add_paragraph("Публикации" + (" за указанный период" if start_date or end_date else "") + " не найдены.")
    else:
        headers = ["№\nп/п", "Наименование работы", "Форма\nработы", "Выходные данные", "Объем в\nп.л.", "Соавторы"]
        num_columns = len(headers)
        table = document.add_table(rows=1, cols=num_columns); table.style = 'Table Grid'; table.autofit = False; table.allow_autofit = False; table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # -- Установка ширины колонок (Форма работы = 2.8) --
        try:
            widths_cm = [1.0, 5.5, 2.8, 4.7, 1.5, 2.5]
            if len(widths_cm) == num_columns:
                 for i, width in enumerate(widths_cm): table.columns[i].width = Cm(width)
                 for i, width in enumerate(widths_cm): table.rows[0].cells[i].width = Cm(width) # Применяем и к хедеру
            else: logger.warning("Ширины не совпадают с кол-вом столбцов.")
        except Exception as e: logger.warning(f"Ошибка установки ширины колонок: {e}")
        # -- Конец установки ширины --

        # Заполнение заголовков
        hdr_cells = table.rows[0].cells;
        for i, header_text in enumerate(headers):
            p = hdr_cells[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; parts = header_text.split('\n')
            run = p.add_run(parts[0]); run.font.bold = True; run.font.size = Pt(11);
            if len(parts) > 1:
                for part in parts[1:]: run = p.add_run('\n' + part); run.font.bold = True; run.font.size = Pt(11);
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Строка нумерации
        num_row_cells = table.add_row().cells
        for i, cell in enumerate(num_row_cells):
             p = cell.paragraphs[0]; run = p.add_run(str(i + 1)); run.font.name = 'Times New Roman'; run.font.size = Pt(11); run.font.bold = True
             p.alignment = WD_ALIGN_PARAGRAPH.CENTER; cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
             if 'widths_cm' in locals() and len(widths_cm) == num_columns: cell.width = Cm(widths_cm[i])

        # --- Функция для добавления строк данных V5 ---
        def add_data_row(publication, index, group_char):
            row_cells = table.add_row().cells
            # Настройка ячеек и шрифта
            default_align=WD_ALIGN_PARAGRAPH.LEFT; default_valign=WD_ALIGN_VERTICAL.TOP
            center_align=WD_ALIGN_PARAGRAPH.CENTER; center_valign=WD_ALIGN_VERTICAL.CENTER
            if 'widths_cm' in locals() and len(widths_cm) == num_columns:
                for i, width in enumerate(widths_cm): row_cells[i].width = Cm(width)
            for cell in row_cells: cell.vertical_alignment = default_valign
            for p in cell.paragraphs: p.alignment = default_align; p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
            base_font = {'name': 'Times New Roman', 'size': Pt(11)}
            def set_font(run): run.font.name = base_font['name']; run.font.size = base_font['size']; return run

            # 1. № п/п
            p = row_cells[0].paragraphs[0]; set_font(p.add_run(str(index))); p.alignment = center_align; row_cells[0].vertical_alignment = center_valign;

            # 2. Наименование работы (Тип С БОЛЬШОЙ БУКВЫ, ПЭВМ - КАПСОМ)
            p = row_cells[1].paragraphs[0]
            type_suffix = ""
            pub_type_name = publication.type.name if publication.type else None
            pub_display_name = publication.display_name.display_name if publication.display_name else None
            type_map_capitalized = {'article': 'Статья', 'conference': 'Конференция', 'misc': 'Прочее', 'book': 'Книга'} # 'ppc' обрабатывается отдельно

            # Явная проверка на ПЭВМ для капса
            if pub_type_name == pvm_type_name:
                type_suffix = " (ПЭВМ)" # Прямое указание капса
            elif pub_display_name:
                 # Отображаем display_name как есть, если он не содержит индексы ВАК/WoS/Scopus/РИНЦ или если это один из спец.типов
                 is_simple_display = all(idx not in pub_display_name for idx in ["ВАК", "WoS", "Scopus", "РИНЦ"])
                 is_special_type_name = pub_display_name in [uchebnoe_posobie_display_name, monograph_display_name] # Проверяем по именам
                 if is_simple_display or is_special_type_name:
                     type_suffix = f" ({pub_display_name})" # Используем имя как есть (из БД ожидается с большой буквы)
                 # Если display_name сложный и не ПЭВМ/Моно/Пособие - используем маппинг по type.name
                 elif pub_type_name in type_map_capitalized:
                     type_suffix = f" ({type_map_capitalized[pub_type_name]})"
                 # Фолбэк на capitalize внутреннего имени, если тип не в маппинге
                 elif pub_type_name:
                     type_suffix = f" ({pub_type_name.capitalize()})"
            # Если нет display_name, используем маппинг или capitalize для type.name
            elif pub_type_name in type_map_capitalized:
                type_suffix = f" ({type_map_capitalized[pub_type_name]})"
            elif pub_type_name:
                type_suffix = f" ({pub_type_name.capitalize()})"

            set_font(p.add_run(f"{publication.title or 'Без названия'}{type_suffix}"));

            # 3. Форма работы
            p = row_cells[2].paragraphs[0]; set_font(p.add_run(publication.work_form or "Печатная"));

            # 4. Выходные данные (С ЛОГИКОЙ ДЛЯ КАЖДОГО БЛОКА V5)
            output_parts = []
            p = row_cells[3].paragraphs[0]

            # --- Логика для блока б) ПЭВМ ---
            if group_char == 'б':
                notes_text = (publication.notes or "").strip()
                run = set_font(p.add_run(notes_text if notes_text else "-"))
                if '\n' in notes_text: p.paragraph_format.line_spacing = 1.0
            # --- Логика для блока а) ВСЕ ОСТАЛЬНЫЕ ---
            elif group_char == 'а':
                # Наименование журнала/конференции (если есть)
                if publication.journal_conference_name:
                    output_parts.append(publication.journal_conference_name)
                # Место издательства и Издательство - ВСЕГДА, если есть
                if publication.publisher_location:
                    output_parts.append(publication.publisher_location)
                if publication.publisher:
                    output_parts.append(publication.publisher)
                # Год
                if publication.year: output_parts.append(str(publication.year))
                # Том, Номер, Страницы
                if publication.volume: output_parts.append(f"Т.{publication.volume}")
                if publication.number: output_parts.append(f"№{publication.number}")
                if publication.pages:
                    # Формат страниц зависит от того, Монография ли это (в группе А могут быть только монографии типа book)
                    is_monograph_a = pub_type_name == 'book' and pub_display_name == monograph_display_name
                    if is_monograph_a: # Формат "X c."
                        pages_clean = "".join(filter(str.isdigit, publication.pages.split('-')[0]))
                        if pages_clean: output_parts.append(f"{pages_clean} c.")
                    else: # Формат "С.X-Y"
                        output_parts.append(f"С.{publication.pages.replace('-', '–')}")
                # Запись в ячейку
                run = set_font(p.add_run(", ".join(filter(None, output_parts))));
            # --- Логика для блока в) УЧЕБНОЕ ПОСОБИЕ ---
            elif group_char == 'в':
                # Место издательства, Издательство
                if publication.publisher_location: output_parts.append(publication.publisher_location)
                if publication.publisher: output_parts.append(publication.publisher)
                # Год
                if publication.year: output_parts.append(str(publication.year))
                # Том, Номер
                if publication.volume: output_parts.append(f"Т.{publication.volume}")
                if publication.number: output_parts.append(f"№{publication.number}")
                # Страницы (формат "X c.")
                if publication.pages:
                    pages_clean = "".join(filter(str.isdigit, publication.pages.split('-')[0]))
                    if pages_clean: output_parts.append(f"{pages_clean} c.")
                # Запись в ячейку
                run = set_font(p.add_run(", ".join(filter(None, output_parts))));

            # 5. Объем в п.л. (без изменений)
            volume_str = "-"
            volume = publication.printed_sheets_volume
            if volume is not None and volume > 0:
                try: volume_float=float(volume); volume_str = str(int(volume_float)) if volume_float==int(volume_float) else f"{volume_float:g}".replace('.', ',')
                except: volume_str = "?"
            p = row_cells[4].paragraphs[0]; set_font(p.add_run(volume_str)); p.alignment = center_align; row_cells[4].vertical_alignment = center_valign;

            # 6. Соавторы (без изменений)
            coauthors_list = []
            if publication.authors: author_names = [a.name.strip() for a in publication.authors if a.name]; coauthors_list = [name for name in author_names if name.lower() != user.full_name.strip().lower()]
            p = row_cells[5].paragraphs[0]; set_font(p.add_run(", ".join(coauthors_list) if coauthors_list else "-"));
        # -- Конец функции add_data_row --

        # Заполнение данными по группам
        global_index = 0
        if group_a: _add_section_header(table, 'а)', num_columns); [add_data_row(pub, i + 1, 'а') for i, pub in enumerate(group_a, global_index)]; global_index += len(group_a)
        if group_b: _add_section_header(table, 'б)', num_columns); [add_data_row(pub, i + 1, 'б') for i, pub in enumerate(group_b, global_index)]; global_index += len(group_b)
        if group_c: _add_section_header(table, 'в)', num_columns); [add_data_row(pub, i + 1, 'в') for i, pub in enumerate(group_c, global_index)]; global_index += len(group_c)


    # 6. Подписи
    document.add_paragraph(); p_prepod = document.add_paragraph(); p_prepod.paragraph_format.space_before = Pt(12); p_prepod.paragraph_format.line_spacing = 1.0
    p_prepod.add_run("Подпись преподавателя" + "\t\t" + "____________________" + "\t" + f"{user_position}, {user.full_name}")
    p_zavkaf = document.add_paragraph(); p_zavkaf.paragraph_format.line_spacing = 1.0; zavkaf_position = "заведующего кафедрой"; zavkaf_name = "(ФИО Заведующего)"; # Заглушка
    p_zavkaf.add_run(f"Подпись {zavkaf_position}" + "\t" + "____________________" + "\t" + f"{zavkaf_name}")

    # 7. Вторая таблица (шаблон)
    document.add_paragraph(); headers2 = ["№\nп/п", "Наименование дисциплины", "Электронная\nсреда", "Соавторы", "Экспертное\nзаключение"]; num_columns2 = len(headers2)
    table2 = document.add_table(rows=1, cols=num_columns2); table2.style = 'Table Grid'; table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells2 = table2.rows[0].cells
    for i, header_text in enumerate(headers2):
         p = hdr_cells2[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; parts = header_text.split('\n')
         run = p.add_run(parts[0]); run.font.bold = True; run.font.size = Pt(11); run.font.name = 'Times New Roman'
         if len(parts) > 1:
             for part in parts[1:]: run = p.add_run('\n' + part); run.font.bold = True; run.font.size = Pt(11); run.font.name = 'Times New Roman'
         hdr_cells2[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # 8. Сохранение в байтовый поток
    buffer = BytesIO();
    try:
        document.save(buffer); buffer.seek(0); logger.info("Генерация DOCX отчета (ПМИ v5) завершена.")
    except Exception as save_err:
        logger.error(f"Ошибка сохранения DOCX: {save_err}", exc_info=True); document_error = Document(); document_error.add_paragraph(f"Ошибка сервера при сохранении документа: {save_err}")
        buffer_error = BytesIO(); document_error.save(buffer_error); buffer_error.seek(0); return buffer_error.getvalue()
    return buffer.getvalue()